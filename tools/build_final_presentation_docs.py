from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
SPEAKER_PLAN_PATH = OUTPUTS / "Dolgin_Wenstrom_Final_Project_Speaker_Plan.docx"
SCRIPT_PATH = OUTPUTS / "Dolgin_Wenstrom_Final_Project_Presentation_Script.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "0A7F98"
AMBER = "9A6500"
MUTED = "5E6B78"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
WHITE = "FFFFFF"
BORDER = "C8D2DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color=BORDER, size=8, space=5, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_shading(paragraph, fill=LIGHT_FILL):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom_styles = {
        "Project Kicker": (10.5, CYAN, True, False, 0, 3, 1.0),
        "Project Title": (29, INK, True, False, 0, 8, 1.0),
        "Project Subtitle": (13.2, MUTED, False, False, 0, 18, 1.12),
        "Timing Line": (9.5, MUTED, True, False, 0, 8, 1.0),
        "Speaker Cue": (11.5, BLUE, True, False, 8, 3, 1.0),
        "Stage Direction": (9.5, MUTED, False, True, 3, 7, 1.12),
        "Closing Cue": (10, AMBER, True, False, 5, 6, 1.0),
    }
    for name, (size, color, bold, italic, before, after, line) in custom_styles.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, 1)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line


def add_numbering(doc, abstract_id, num_id, kind):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_numbering(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_run_font(lead, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)
    return paragraph


def configure_document(doc, running_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(running_label)
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    label = footer_p.add_run("Beyond Points Per Game  |  ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(footer_p)

    configure_styles(doc)
    add_numbering(doc, 42, 42, "bullet")
    add_numbering(doc, 43, 43, "decimal")

    core = doc.core_properties
    core.title = running_label
    core.subject = "DATASCI 209 final project presentation"
    core.author = "Kenny Dolgin and Ryan Wenstrom"


def add_title_block(doc, title, subtitle, metrics):
    kicker = doc.add_paragraph("DATASCI 209  |  FINAL PROJECT", style="Project Kicker")
    kicker.paragraph_format.keep_with_next = True
    title_p = doc.add_paragraph(title, style="Project Title")
    title_p.paragraph_format.keep_with_next = True
    subtitle_p = doc.add_paragraph(subtitle, style="Project Subtitle")
    subtitle_p.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [9360 // len(metrics)] * len(metrics)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths, indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, TABLE_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        value_run = p.add_run(value + "\n")
        set_run_font(value_run, size=13, color=INK, bold=True)
        label_run = p.add_run(label.upper())
        set_run_font(label_run, size=8, color=MUTED, bold=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, LIGHT_FILL)
    set_paragraph_border(p, color=CYAN, size=16, space=8, side="left")
    lead = p.add_run(label.upper() + "  ")
    set_run_font(lead, size=9.5, color=CYAN, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=INK)
    return p


def add_table_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def build_speaker_plan():
    doc = Document()
    configure_document(doc, "Final Presentation Speaker Plan")
    add_title_block(
        doc,
        "Final Presentation Speaker Plan",
        "A 15-minute, two-presenter run of show for Beyond Points Per Game",
        [("15:00", "total"), ("10", "slides"), ("2", "live demos"), ("6", "testers")],
    )
    add_callout(
        doc,
        "Presentation objective",
        "Show how the design changed after usability testing, demonstrate the final visualization in action, explain the evidence behind those changes, and justify the final tool choices. The midterm appears only as a brief before-and-after reference.",
    )

    doc.add_paragraph("Run of show", style="Heading 1")
    headers = ["Slide", "Lead", "Time", "Purpose and on-screen action", "Handoff"]
    rows = [
        ("1", "Kenny", "0:30", "Open with the final thesis and frame the talk around change, testing, and the final interactive.", "Continue into the delta."),
        ("2", "Kenny", "1:00", "Establish the midterm-to-final delta: neutral framing, persistent state, exact values, and a roster-level workflow.", "Move directly to the live Role Explorer."),
        ("3", "Kenny", "3:00", "Live Role Explorer: select LeBron 2020 LAL and Harden 2025 LAC; show the insight, role map, fingerprint, and similar seasons.", "Ryan opens Team Roster Lab."),
        ("4", "Ryan", "3:00", "Live Team Roster Lab: choose Oklahoma City, inspect one player, Coverage, Benchmark, add Boston, then use Swap teams.", "Stay on the deck for the study."),
        ("5", "Ryan", "1:30", "Explain one moderated round with six participants and what behavior and interpretation evidence was recorded.", "Continue directly to findings."),
        ("6", "Ryan", "1:30", "Lead with the recurring interpretation problem: ranking and causation risk, unclear terms, dense evidence, and lost hover state.", "Continue into the design response."),
        ("7", "Ryan", "1:30", "Tie point-of-use definitions, neutral titles, persistent selection, exact values, and repeated guardrails to observed failures.", "Hand the larger redesign to Kenny."),
        ("8", "Kenny", "1:15", "Explain why the geography view was removed and how the Roster Lab better answers the roster-responsibility question.", "Continue to tool choices."),
        ("9", "Kenny", "0:45", "Justify Flask/Jinja, D3/JavaScript, CSV/JSON, the static build, and GitHub Pages.", "Continue to the conclusion."),
        ("10", "Kenny", "1:00", "Resolve the thesis, name limitations, state what testing improved, and invite questions.", "End on the live URL."),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [620, 1240, 740, 5130, 1630]
    set_table_geometry(table, widths, indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TABLE_FILL)
        add_table_text(cell, header, bold=True, color=INK, size=9)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row_data)):
            add_table_text(
                cell,
                value,
                bold=idx in (0, 1, 2),
                color=BLUE if idx == 1 else INK,
                size=8.7 if idx in (3, 4) else 9,
                align=WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT,
            )
        if int(row_data[0]) % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
    set_table_geometry(table, widths, indent_dxa=120)

    doc.add_paragraph("Live-demo operating plan", style="Heading 1")
    doc.add_paragraph("Role Explorer - Kenny (Slide 3)", style="Heading 2")
    role_steps = [
        "Open the Role Explorer and state the question: how do two player-season profiles differ without turning the view into a ranking?",
        "Set Player A to LeBron James, 2020 LAL, and Player B to James Harden, 2025 LAC.",
        "Read the comparison insight, then use the role map and cards to explain profile shape.",
        "Scroll to the Role fingerprint. Point out that percentiles are relative ranks in the reviewed dataset and that exact values remain available.",
        "Show the closest comparable seasons and say: statistical similarity describes profile shape, not player quality.",
    ]
    for item in role_steps:
        add_list_item(doc, item, 43)

    doc.add_paragraph("Team Roster Lab - Ryan (Slide 4)", style="Heading 2")
    team_steps = [
        "Open Team rosters, keep the current season, and choose Oklahoma City.",
        "Select one player and use the player tray to connect the roster map to exact metrics.",
        "Open Coverage to show raw values beside season-league percentiles.",
        "Open Benchmark. In Compare team, choose Boston and select Apply.",
        "Use Swap selected and comparison. Confirm that Boston becomes the active roster and Oklahoma City remains the comparison.",
        "State the limitation: the top 10 by minutes are a rotation proxy, not an official depth chart, and the comparison is descriptive rather than causal.",
    ]
    for item in team_steps:
        add_list_item(doc, item, 43)

    doc.add_paragraph("Presenter ownership and handoffs", style="Heading 1")
    add_callout(
        doc,
        "Kenny's role",
        "Own Slides 1-3 and 8-10: the opening thesis, midterm-to-final delta, Role Explorer demo, geography-to-Roster-Lab redesign, tool choices, and conclusion.",
    )
    add_callout(
        doc,
        "Ryan's role",
        "Own Slides 4-7: the Team Roster Lab demo, usability-study method, cross-participant findings, and the point-of-use design response.",
    )
    handoffs = [
        "Kenny to Ryan after Slide 3: 'That is the individual-role workflow. Ryan will show how the same responsibility question scales up to a full team rotation.'",
        "Ryan to Kenny after Slide 7: 'Those point-of-use changes addressed the immediate failures. Kenny will show why we also replaced the entire geography workflow.'",
    ]
    for item in handoffs:
        add_list_item(doc, item, 42)

    doc.add_paragraph("Requirement coverage", style="Heading 1")
    requirements = [
        "What changed and why: Slides 2, 7, and 8.",
        "At least one usability round with at least three subjects: Slide 5 documents one round with six participants.",
        "Testing conclusions: Slide 6.",
        "Design changes in response to findings: Slides 7 and 8.",
        "Tool choices: Slide 9.",
        "All team members participate: each presenter leads one live demo and multiple evidence slides.",
    ]
    for item in requirements:
        add_list_item(doc, item, 42)

    doc.add_paragraph("Rehearsal checkpoints", style="Heading 1")
    checkpoints = [
        "Rehearse from a fresh browser tab and verify both named player-seasons are selectable.",
        "Preload the Role Explorer and Team Roster Lab before class; keep Slides 3 and 4 available as large fallback images.",
        "Do not narrate every chart. Use the dashboard only to answer the two demo questions.",
        "Call time at the end of Slides 4, 6, and 8. Target 6:30, 9:30, and 12:15 elapsed.",
        "If a demo stalls, return to the deck immediately and use the fallback screenshot while keeping the same explanation.",
        "End by 15:00, including the limitations and the invitation for questions.",
    ]
    for item in checkpoints:
        add_list_item(doc, item, 42)

    OUTPUTS.mkdir(parents=True, exist_ok=True)
    doc.save(SPEAKER_PLAN_PATH)


def add_script_heading(doc, slide, title, presenter, timing, mode):
    doc.add_paragraph(f"Slide {slide}: {title}", style="Heading 1")
    meta = doc.add_paragraph(style="Timing Line")
    run = meta.add_run(f"{presenter}  |  {timing}  |  {mode}")
    set_run_font(run, size=9.5, color=MUTED, bold=True)


def add_speaker_paragraph(doc, speaker, text):
    cue = doc.add_paragraph(f"{speaker}", style="Speaker Cue")
    cue.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_action(doc, text):
    p = doc.add_paragraph(style="Stage Direction")
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    set_paragraph_shading(p, LIGHT_FILL)
    set_paragraph_border(p, color=CYAN, size=12, space=6, side="left")
    lead = p.add_run("ACTION  ")
    set_run_font(lead, size=9, color=CYAN, bold=True, italic=False)
    body = p.add_run(text)
    set_run_font(body, size=9.5, color=MUTED, italic=True)


def add_handoff(doc, text):
    p = doc.add_paragraph(style="Closing Cue")
    set_paragraph_border(p, color=AMBER, size=10, space=5, side="left")
    lead = p.add_run("HANDOFF  ")
    set_run_font(lead, size=9, color=AMBER, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10, color=INK, bold=True)


def add_slide_break(doc):
    doc.add_page_break()


def build_script():
    doc = Document()
    configure_document(doc, "Final Presentation Script")
    add_title_block(
        doc,
        "Final Presentation Script",
        "Beyond Points Per Game - final dashboard demonstration and usability readout",
        [("15:00", "runtime"), ("10", "slides"), ("2", "presenters"), ("2", "live demos")],
    )
    add_callout(
        doc,
        "How to use this script",
        "Speak conversationally rather than reading every word. Stage directions are operational cues, not spoken lines. The script is paced for 15 minutes and deliberately emphasizes the final design, usability evidence, and design response instead of repeating the midterm.",
    )
    add_callout(
        doc,
        "Core thesis",
        "Points per game is familiar, but it does not explain how responsibility is distributed. The final project helps people compare player-season role shapes and team rotations while keeping winning as descriptive context rather than a ranking or causal formula.",
    )
    doc.add_paragraph("Timing checkpoints", style="Heading 1")
    for item in [
        "End Slide 2 at 1:30 elapsed.",
        "End both live demonstrations at 7:30 elapsed.",
        "End usability findings at 10:30 elapsed.",
        "End design response at 13:15 elapsed.",
        "Close by 15:00.",
    ]:
        add_list_item(doc, item, 42)
    add_slide_break(doc)

    add_script_heading(doc, 1, "Beyond Points Per Game", "Kenny", "0:30", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Our project starts with a simple problem: points per game is the statistic almost everyone recognizes, but it does not tell us how a player creates value or how responsibility is distributed across a roster.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Today we are not going to replay the midterm. We will show what changed, demonstrate the final player and team workflows, and explain how one usability-testing round with six people changed the design.",
    )
    add_action(doc, "Kenny continues directly into the midterm-to-final delta.")
    add_slide_break(doc)

    add_script_heading(doc, 2, "What Changed", "Kenny", "1:00", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "At the midterm, the project was mostly a set of player-level exploratory charts. The framing sometimes sounded evaluative - as if a profile could look more like winning basketball - and important detail depended on hover. The final version is a guided comparison workflow. Selections stay visible, exact values sit beside relative measures, and the language repeatedly separates statistical shape from player quality or causal impact.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "We also changed the unit of the second workflow. Instead of asking where teams are located, the Team Roster Lab asks how one team-season distributes scoring, creation, rebounding, efficiency, defensive events, and role breadth. That change matters because it makes the team view answer the same responsibility question as the player view.",
    )
    add_action(doc, "Use the before-and-after images only to establish the delta. Do not explain the old charts in detail.")
    add_action(doc, "Say: 'The fastest way to see the difference is in the final dashboard.' Switch to the live Role Explorer.")
    add_slide_break(doc)

    add_script_heading(doc, 3, "Live Demo - Role Explorer", "Kenny", "3:00", "Live dashboard")
    add_action(doc, "Open the Role Explorer. Set Player A to LeBron James, 2020 LAL, and Player B to James Harden, 2025 LAC.")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The live question is not who is better. It is: how do these two player-seasons differ in role shape, and what historical team-win context surrounds statistically similar seasons?",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The controls make the player-season the unit explicit. That was important because a tester initially read one mark as one game. Once both seasons are selected, the comparison sentence summarizes the evidence without turning it into a verdict.",
    )
    add_action(doc, "Point to both player cards and the comparison insight. Then scroll to the role map.")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The role map positions the two selected seasons within the reviewed set. The selected marks and cards stay visible, so the user does not lose context while scrolling. Nearby seasons are similar in the variables used here, but proximity is not a quality score.",
    )
    add_action(doc, "Scroll to Role fingerprint. Point to one metric where a bar is above the 90th percentile and to the aligned value pill outside the bar.")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The fingerprint compares scoring, creation, rebounding, defensive events, shooting efficiency, shot volume, turnovers, and role breadth as league percentiles. The labels now sit in a consistent value column, including above the 90th percentile, so they no longer disappear into long bars. Hover still provides the actual metric value, but the comparison does not depend on hover alone.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The note directly above the evidence also matters: higher is descriptive, not always better. A high turnover or shot-volume percentile can reflect responsibility and workload, not simply a positive or negative grade.",
    )
    add_action(doc, "Show closest comparable seasons and the similarity explanation.")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Finally, closest matches answer a profile-similarity question. The distance uses the listed percentile dimensions. A small distance means a similar statistical shape; it does not mean similar quality, salary value, or causal impact on winning.",
    )
    add_handoff(doc, "Say: 'That is the individual-role workflow. Ryan will show how the same responsibility question scales up to a team rotation.'")
    add_slide_break(doc)

    add_script_heading(doc, 4, "Live Demo - Team Roster Lab", "Ryan", "3:00", "Live dashboard")
    add_action(doc, "Open Team rosters. Keep the current season and select Oklahoma City.")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The Team Roster Lab treats the team-season as the unit. The rail is sorted by win percentage for orientation, but the main question is how responsibility is distributed within the selected roster.",
    )
    add_action(doc, "Select one player in the roster map and point to the persistent player tray.")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The roster map uses scoring load horizontally, non-scoring contribution vertically, and minutes for bubble size. Selecting a player updates a persistent tray with raw values and percentiles, so an exact answer does not require holding a hover state.",
    )
    add_action(doc, "Open Coverage. Point to one row and compare a raw value with its percentile.")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Coverage turns the same evidence into a scan-friendly table. Participants often preferred bars and tables for precise questions, so this view keeps raw values and season-league percentiles together.",
    )
    add_action(doc, "Open Benchmark. Open Compare team, select Boston, and choose Apply.")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The benchmark compares Oklahoma City with the league midpoint, the average profile of the five highest-win teams, and an optional second team. Adding Boston makes the comparison explicit without replacing the active roster.",
    )
    add_action(doc, "Open Compare: BOS and choose Swap selected and comparison. Confirm that Boston becomes selected and Oklahoma City remains the comparison.")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Swap teams turns the comparison into a two-way exploration. Boston is now the active roster across the summary, roster map, player tray, and coverage table, while Oklahoma City stays in the benchmark as the comparison. The URL updates too, so this state can be shared or restored.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The limitation stays visible: the top 10 players by minutes are a rotation proxy, not an official depth chart, and these roster shapes are descriptive context. They do not prove that one construction causes winning.",
    )
    add_action(doc, "Return to Slide 5. Ryan continues: 'Those interaction choices came from one moderated usability round.'")
    add_slide_break(doc)

    add_script_heading(doc, 5, "Usability Study", "Ryan", "1:30", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "We ran one moderated round with six different participants. I conducted sessions with Ryan P01 through P03, and Kenny conducted K-01 Paul, K-02 Caroline, and K-03 Sean. The sessions were planned for about 30 minutes each and included people with different levels of NBA familiarity and chart literacy.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Participants first oriented themselves without coaching. Then they compared player-seasons, explained similarity and percentiles, and used the team and methods views. We recorded first action, evidence used, hesitation, help needed, confidence, and interpretation errors. That distinction mattered: a control can be easy to click while the resulting conclusion is still wrong.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The varied expertise exposed failures that more experienced dashboard users could work around. It also gave us repeated evidence across participants rather than a single preference.",
    )
    add_action(doc, "Ryan continues: 'The interaction was usually manageable. The recurring problem was interpretation.'")
    add_slide_break(doc)

    add_script_heading(doc, 6, "What We Learned", "Ryan", "1:30", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "The central finding was that interpretation, not basic interaction, was the recurring risk. Ryan P03 and Kenny K-03 explicitly read parts of the prototype as a ranking or a recipe for winning. Other participants independently flagged wording that felt evaluative.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Definitions were the next repeated problem. Participants paused on eFG percentage, per-36 statistics, percentiles, share denominators, defensive events, comparable-season median, and similarity distance. The information sometimes existed, but it was too far from the moment when the user needed it.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Dense charts also made exact tasks harder. Some participants focused on famous players or isolated marks rather than the distribution, while bars and tables made precise comparisons easier. Finally, hover-only state was easy to lose during scrolling. Those four patterns became the priorities: neutral framing, local definitions, exact values, and persistent state.",
    )
    add_action(doc, "Ryan continues: 'Those priorities led directly to the visible changes in the final interface.'")
    add_slide_break(doc)

    add_script_heading(doc, 7, "Design Response - Point of Use", "Ryan", "1:30", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Ryan",
        "We treated each visible change as a response to observed behavior. Ranking risk led to neutral questions and repeated statements that similarity is not quality. Causal language led us to keep winning as context and repeat that the evidence is descriptive.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Terminology problems led to point-of-use definitions. Instead of expecting users to leave the chart for a methods page, short explanations now appear where percentiles, per-36 measures, defensive events, and similarity are introduced. The methods page remains available for the full calculation.",
    )
    add_speaker_paragraph(
        doc,
        "Ryan",
        "Dense evidence and lost hover state led to persistent player cards, selected marks, raw values beside percentiles, and table alternatives. The final fingerprint is a good example: the guardrail and exact-value access live with the evidence, and every percentile label now sits in the same right-side gutter so long bars cannot trap the 90th-to-100th-percentile labels.",
    )
    add_handoff(doc, "Say: 'Those point-of-use changes addressed the immediate failures. Kenny will show why we also replaced the entire geography workflow.'")
    add_slide_break(doc)

    add_script_heading(doc, 8, "Design Response - From Geography to Rosters", "Kenny", "1:15", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The earlier Team Context Map was mechanically understandable, but geography did not advance the player-role question. State-level averages could also hide meaningful differences between teams that happen to share a state.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "We replaced it with the Team Roster Lab. The unit is now a team-season rotation, and the role map, coverage matrix, and benchmark all answer one coherent question: how is responsibility distributed across this roster? Swap teams also makes the comparison reversible without losing the prior team. The limitations remain visible: the top 10 by minutes are a proxy, and we do not claim to model lineups, injuries, coaching, tactics, or complete defense.",
    )
    add_action(doc, "Kenny continues: 'Once the interaction matched the question, the final stack had to keep it custom and dependable.'")
    add_slide_break(doc)

    add_script_heading(doc, 9, "Tool Choice", "Kenny", "0:45", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Flask and Jinja kept local routes, templates, and data endpoints simple while we were iterating. D3 with vanilla JavaScript gave us direct control over linked selections, custom SVG charts, and URL-backed state without adding a framework layer. CSV and JSON keep the data path inspectable. For the public version, a static build on GitHub Pages removes server-hosting complexity while preserving the interactive behavior.",
    )
    add_action(doc, "Kenny continues directly into the conclusion.")
    add_slide_break(doc)

    add_script_heading(doc, 10, "Conclusion", "Kenny", "1:00", "Slide on screen")
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The final dashboard is no longer a collection of charts that asks the user to infer the reading path. It is a guided comparison workflow with persistent state, exact values, and a roster-level view that answers the same responsibility question at a second scale.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Usability testing showed that plain language and local definitions build trust, bars and tables support precise answers, and guardrails have to be repeated where overinterpretation can happen. The project remains descriptive: it is not a causal model, a player grade, an official depth chart, or a complete model of defense and team context.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "The live project is available at the URL on screen.",
    )
    add_speaker_paragraph(
        doc,
        "Kenny",
        "Thank you. We are happy to take questions.",
    )

    OUTPUTS.mkdir(parents=True, exist_ok=True)
    doc.save(SCRIPT_PATH)


if __name__ == "__main__":
    build_speaker_plan()
    build_script()
    print(SPEAKER_PLAN_PATH)
    print(SCRIPT_PATH)
